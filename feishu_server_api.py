import requests
import json
from datetime import datetime, timedelta
from pathlib import Path
import re
from docx import Document


class FeishuDocsAPI:
    def __init__(
        self,
        notice_message_heading,
        notice_dir,
        document_id,
        app_id,
        app_secret,
        message_heading_text,
        debug_offline_all_document_comments_response_json=False,
        save_all_document_blocks_response_as_json=True,
        save_all_document_comments_response_as_json=True,
        debug_offline_all_document_blocks_response_json=False,
    ) -> None:
        self.notice_dir = Path(notice_dir)
        self.notice_path = Path()
        script_dir = Path(__file__).resolve().parent
        self.notice_message_heading = notice_message_heading
        self.document_id = document_id
        tenant_access_token = self.obtain_tenant_access_token(app_id, app_secret)
        self.access_token = tenant_access_token
        # Process all document comments firstly.
        # Unsolved comments need to be updated to the document.
        self.all_document_comments_response = None
        self.unsolved_document_comments_list = []
        self.preprocess_all_document_comments(
            debug_offline_all_document_comments_response_json
        )
        all_document_comments_response_path = (
            script_dir / "all_document_comments_response.json"
        )
        if save_all_document_comments_response_as_json is True:
            with open(all_document_comments_response_path, "w+", encoding="utf8") as f:
                json.dump(self.all_document_comments_response, f, ensure_ascii=False)
        # Process all document blocks.
        self.all_document_blocks_response = None
        self.all_document_blocks = []
        # The first block is the document title.
        self.all_document_children_block_ids = []
        self.message_heading_text = message_heading_text
        self.message_heading_block_id = None
        self.item_message_heading_block_index = None
        self.children_message_heading_block_index = None
        self.item_message_block_start_index = None
        self.children_message_block_start_index = None
        self.message_blocks_list = []
        self.preprocess_all_document_blocks(
            debug_offline_all_document_blocks_response_json
        )
        all_document_blocks_response_path = (
            script_dir / "all_document_blocks_response.json"
        )
        if save_all_document_blocks_response_as_json is True:
            with open(all_document_blocks_response_path, "w+", encoding="utf8") as f:
                json.dump(self.all_document_blocks_response, f, ensure_ascii=False)

    def obtain_tenant_access_token(self, app_id, app_secret):
        url = "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal"
        headers = {
            "Content-Type": "application/json; charset=utf-8",
        }
        payload = json.dumps(
            {
                "app_id": app_id,
                "app_secret": app_secret,
            }
        )
        response = requests.request("POST", url, headers=headers, data=payload)
        response_json = json.loads(response.text)
        return response_json["tenant_access_token"]

    # Methods for comments.

    def obtain_all_document_comments(self):
        url = (
            "https://open.feishu.cn/open-apis/drive/v1/files/"
            + self.document_id
            + "/comments?file_type=docx"
        )
        payload = ""
        access_token = "Bearer " + self.access_token
        headers = {"Authorization": access_token}
        response = requests.request("GET", url, headers=headers, data=payload)
        # print(response.text)
        response_json = json.loads(response.text)
        return response_json

    def delete_a_reply(self, document_comment):
        # Deprecated
        # The document app can only delete replies created by itself.
        reply_id = document_comment["reply_list"]["replies"][0]["reply_id"]
        url = (
            "https://open.feishu.cn/open-apis/drive/v1/files/"
            + self.document_id
            + "/comments/"
            + document_comment["comment_id"]
            + "/replies/"
            + reply_id
            + "?file_type=docx"
        )
        payload = ""
        access_token = "Bearer " + self.access_token
        headers = {"Authorization": access_token}
        response = requests.request("DELETE", url, headers=headers, data=payload)
        # print(response.text)
        response_json = json.loads(response.text)
        return response_json

    def solve_a_reply(self, document_comment):
        url = (
            "https://open.feishu.cn/open-apis/drive/v1/files/"
            + self.document_id
            + "/comments/"
            + document_comment["comment_id"]
            + "?file_type=docx"
        )
        payload = json.dumps({"is_solved": True})  # type: ignore
        access_token = "Bearer " + self.access_token
        headers = {
            "Content-Type": "application/json; charset=utf-8",
            "Authorization": access_token,
        }
        response = requests.request("PATCH", url, headers=headers, data=payload)
        # print(response.text)
        response_json = json.loads(response.text)
        return response_json

    def preprocess_all_document_comments(
        self, debug_offline_all_document_comments_json
    ):
        if debug_offline_all_document_comments_json is True:
            script_dir = Path(__file__).resolve().parent
            all_document_comments_response_json_path = (
                script_dir / "all_document_comments_response.json"
            )
            with open(
                all_document_comments_response_json_path, "r", encoding="utf8"
            ) as f:
                self.all_document_comments_response = json.load(f)
        else:
            self.all_document_comments_response = self.obtain_all_document_comments()
        self.all_document_comments = self.all_document_comments_response["data"][
            "items"
        ]
        for document_comment in self.all_document_comments:
            # Process unsolved comments only.
            if document_comment["solver_user_id"] is None:
                self.unsolved_document_comments_list.append(document_comment)
            else:
                continue
        print(
            "{} unsolved document comment(s) obtained".format(
                len(self.unsolved_document_comments_list)
            )
        )
        for unsolved_document_comment in self.unsolved_document_comments_list:
            # Process unsolved comments only.
            reply = unsolved_document_comment["reply_list"]["replies"][0]
            unsolved_document_comment_text = reply["content"]["elements"][0][
                "text_run"
            ]["text"]
            try:
                create_blocks_response_json = self.create_blocks(
                    unsolved_document_comment_text
                )
                if create_blocks_response_json["code"] == 0:
                    create_blocks_response_content = create_blocks_response_json[
                        "data"
                    ]["children"][0]["text"]["elements"][0]["text_run"]["content"]
                    print(
                        "sucessfully created a document children block based on the comment:",
                        create_blocks_response_content,
                    )
                    try:
                        solve_a_reply_response_json = self.solve_a_reply(
                            document_comment
                        )
                        if solve_a_reply_response_json["code"] == 0:
                            print(
                                "sucessfully solve the comment:",
                                unsolved_document_comment_text,
                            )
                    except Exception as e:
                        print(f"A error occurred when calling delete_a_reply: {e}")
            except Exception as e:
                print(f"An error occurred when calling create_blocks: {e}")

    # Methods for document blocks.

    def obtain_plain_document_text_content(self):
        url = (
            "https://open.feishu.cn/open-apis/docx/v1/documents/"
            + self.document_id
            + "/raw_content"
        )
        payload = ""
        access_token = "Bearer " + self.access_token
        headers = {"Authorization": access_token}
        response = requests.request("GET", url, headers=headers, data=payload)
        # print(response.text)
        response_json = json.loads(response.text)
        return response_json

    def obtain_all_document_blocks(self):
        url = (
            "https://open.feishu.cn/open-apis/docx/v1/documents/"
            + self.document_id
            + "/blocks"
        )
        payload = ""
        access_token = "Bearer " + self.access_token
        headers = {"Authorization": access_token}
        response = requests.request("GET", url, headers=headers, data=payload)
        # print(response.text)
        response_json = json.loads(response.text)
        return response_json

    def preprocess_all_document_blocks(
        self, debug_offline_all_document_blocks_response_json
    ):
        if debug_offline_all_document_blocks_response_json is True:
            script_dir = Path(__file__).resolve().parent
            all_document_blocks_response_json_path = (
                script_dir / "all_document_blocks_response.json"
            )
            with open(
                all_document_blocks_response_json_path, "r", encoding="utf8"
            ) as f:
                self.all_document_blocks_response = json.load(f)
        else:
            self.all_document_blocks_response = self.obtain_all_document_blocks()
        self.all_document_blocks = self.all_document_blocks_response["data"]["items"]
        self.all_document_children_block_ids = self.all_document_blocks[0]["children"]
        for block_index, block in enumerate(self.all_document_blocks):
            if (
                block["block_type"] == 3
                and block["heading1"]["elements"][0]["text_run"]["content"]
                == self.message_heading_text
            ):
                self.message_heading_block_id = block["block_id"]
                print(
                    f"{self.message_heading_text} block_id: ",
                    self.message_heading_block_id,
                )
                self.item_message_heading_block_index = block_index
                print(
                    f"{self.message_heading_text} item_block_index: ",
                    self.item_message_heading_block_index,
                )
                # The message_blocks_list does not contain the message title.
                self.message_blocks_list = self.all_document_blocks[
                    self.item_message_heading_block_index + 1 :
                ]
                break
        # The first message block index is the NEXT block of the `message_heading_text``.
        self.item_message_block_start_index = self.item_message_heading_block_index + 1

        # The first block is the document title (parent block) which could not be deleted.
        # Only the children blocks could be deleted.
        # The index is calculated based on children blocks.
        for document_children_block_id_index, document_children_block_id in enumerate(
            self.all_document_children_block_ids
        ):
            if document_children_block_id == self.message_heading_block_id:
                self.children_message_heading_block_index = (
                    document_children_block_id_index
                )
                print(
                    f"{self.message_heading_text} children_block_index: ",
                    self.children_message_heading_block_index,
                )
                self.children_message_block_start_index = (
                    self.children_message_heading_block_index + 1
                )
                break

    def create_blocks(self, text_content):
        url = (
            "https://open.feishu.cn/open-apis/docx/v1/documents/"
            + self.document_id
            + "/blocks/"
            + self.document_id
            + "/children?document_revision_id=-1"
        )
        payload = json.dumps(
            {
                "children": [
                    {
                        "block_type": 2,
                        "text": {
                            "elements": [
                                {
                                    "text_run": {
                                        "content": text_content,
                                        "text_element_style": {},
                                    }
                                },
                            ],
                            "style": {},
                        },
                    }
                ],
                "index": -1,
            }
        )
        access_token = "Bearer " + self.access_token
        headers = {
            "Authorization": access_token,
            "Content-Type": "application/json; charset=utf-8",
        }
        response = requests.request("POST", url, headers=headers, data=payload)
        response_json = json.loads(response.text)
        return response_json

    def update_blocks(self, block):
        block_id = block["block_id"]
        url = (
            "https://open.feishu.cn/open-apis/docx/v1/documents/"
            + self.document_id
            + "/blocks/"
            + block_id
        )
        now = datetime.now()
        formatted_time = now.strftime("%Y-%m-%d %H:%M:%S")
        # print(f"formatted_time: {formatted_time}")
        block_element = block["text"]["elements"][0]
        block_element["text_run"]["content"] = (
            formatted_time + "【已通知】" + block_element["text_run"]["content"]
        )
        payload = json.dumps({"update_text_elements": {"elements": [block_element]}})

        access_token = "Bearer " + self.access_token
        headers = {
            "Authorization": access_token,
            "Content-Type": "application/json; charset=utf-8",
        }

        response = requests.request("PATCH", url, headers=headers, data=payload)
        if response.status_code == 200:
            response_json_data = response.json()
            response_json_element = response_json_data["data"]["block"]["text"][
                "elements"
            ][0]
            response_json_content = response_json_element["text_run"]["content"]
            print("updated block:", response_json_content)

    def check_notice_exists(self):
        now = datetime.now()
        current_date = now.date()
        # print(f"current_date: {current_date}")
        notice_filename = str(current_date) + " 通知.docx"
        notice_path = self.notice_dir / Path(notice_filename)
        self.notice_path = Path(notice_path)
        if not notice_path.exists():
            notice_template_year_month = current_date.strftime("%Y-%m")
            notice_template_path = (
                self.notice_dir / f"{notice_template_year_month}-xx 通知.docx"
            )
            if notice_template_path.exists():
                notice_path.write_bytes(notice_template_path.read_bytes())

    def deliver_and_reply_messages(self):
        def is_after_yesterday_1900(text_message_notified_time):
            now = datetime.now()
            yesterday = now - timedelta(days=1)
            yesterday_19_00 = datetime(
                yesterday.year, yesterday.month, yesterday.day, 19, 00
            )
            yesterday_23_59 = datetime(
                yesterday.year, yesterday.month, yesterday.day, 23, 59, 59
            )
            return yesterday_19_00 < text_message_notified_time <= yesterday_23_59

        self.check_notice_exists()
        pending_message_blocks_list = []
        for message_block_index, message_block in enumerate(
            self.all_document_blocks[self.item_message_block_start_index :],
            start=self.item_message_block_start_index,
        ):
            if message_block["block_type"] == 2:
                message_block_element = message_block["text"]["elements"][0]
                if (
                    18 < len(message_block_element["text_run"]["content"])
                    and message_block_element["text_run"]["content"][19:24]
                    == "【已通知】"
                ):
                    text_message_notified_time = datetime.strptime(
                        message_block_element["text_run"]["content"][:19],
                        "%Y-%m-%d %H:%M:%S",
                    )
                    if is_after_yesterday_1900(text_message_notified_time):
                        pending_message_blocks_list.append(message_block)
                        print("A message was left after 20:40 yesterday.")
                        continue
                    else:
                        pass
                elif message_block_element["text_run"]["content"] == "":
                    pass
                else:
                    pending_message_blocks_list.append(message_block)

        if len(pending_message_blocks_list) != 0:
            print("len(pending_message_blocks_list):", len(pending_message_blocks_list))
            notice = Document(self.notice_path)
            target_paragraph_text = self.notice_message_heading

            for paragraph_index, paragraph in enumerate(notice.paragraphs):
                if target_paragraph_text in paragraph.text:
                    for pending_message_block_index, pending_message_block in enumerate(
                        pending_message_blocks_list
                    ):
                        # Deliver messages
                        pending_message_block_element = pending_message_block["text"][
                            "elements"
                        ][0]
                        pending_message_element = pending_message_block_element[
                            "text_run"
                        ]["content"]
                        new_paragraph = notice.add_paragraph(pending_message_element)
                        # notice.paragraphs.insert(paragraph_index + 1, new_paragraph)
                        paragraph._p.addnext(new_paragraph._p)
                        notice.save(self.notice_path)
                        print(
                            f"delivered pending_message_element[{pending_message_block_index}]: {pending_message_element}"
                        )

                        # Reply messages
                        self.update_blocks(pending_message_block)

    def delete_document_children_blocks(self, document_children_block_index):
        # The feishu official development document is weird.
        parent_block_id = self.document_id
        url = (
            "https://open.feishu.cn/open-apis/docx/v1/documents/"
            + self.document_id
            + "/blocks/"
            + parent_block_id
            + "/children/batch_delete"
        )
        payload = json.dumps(
            {
                "start_index": document_children_block_index,
                "end_index": document_children_block_index + 1,
            }
        )
        access_token = "Bearer " + self.access_token
        headers = {
            "Authorization": access_token,
            "Content-Type": "application/json; charset=utf-8",
        }
        response = requests.request("DELETE", url, headers=headers, data=payload)
        # print("delete document children blocks response:", response.text)

    def delete_notified_messages(self):
        # Deletion is executed based on children blocks, not item blocks.

        def is_text_message_notified_24_hours_before(message_block_element):
            if (
                18 < len(message_block_element["text_run"]["content"])
                and message_block_element["text_run"]["content"][19:24] == "【已通知】"
            ):
                text_message_notified_time = datetime.strptime(
                    message_block_element["text_run"]["content"][:19],
                    "%Y-%m-%d %H:%M:%S",
                )
                # print(text_message_notified_time)
                current_time = datetime.now()
                time_difference = current_time - text_message_notified_time
                time_difference_seconds = time_difference.total_seconds()
                if 24 * 60 * 60 < time_difference_seconds:
                    return True
                else:
                    return False
            else:
                return False

        def obtain_blank_and_notified_message_blocks():
            deletion_dict = {}
            deletion_images_dict = {}

            # Process different dypes of blocks respectively for debugging.

            # Initialize a dict to record blank message block indexes.
            blank_message_blocks_dict = {}
            # Initialize a dict to record notified text messages block indexes.
            text_message_blocks_dict = {}
            notified_text_message_blocks_dict = {}
            # Initialize a dict to record all image indexes if all text messages are notified.
            image_message_blocks_dic = {}

            # There is no message_heading_text in the `self.message_blocks_list`.
            for children_message_block_index, message_block in enumerate(
                self.message_blocks_list, start=self.children_message_block_start_index
            ):
                if message_block["block_type"] == 2:
                    message_block_element = message_block["text"]["elements"][0]
                    if message_block_element["text_run"]["content"] == "":
                        # Record all blank message blocks
                        blank_message_blocks_dict.update(
                            {children_message_block_index: "blank message blocks"}
                        )
                    else:
                        # Record all text message blocks
                        text_message_blocks_dict.update(
                            {
                                children_message_block_index: message_block_element[
                                    "text_run"
                                ]["content"]
                            }
                        )
                        # Record all notified text message blocks
                        if is_text_message_notified_24_hours_before(
                            message_block_element
                        ):
                            notified_text_message_blocks_dict.update(
                                {
                                    children_message_block_index: message_block_element[
                                        "text_run"
                                    ]["content"]
                                }
                            )
                elif message_block["block_type"] == 27:
                    image_message_blocks_dic.update(
                        {children_message_block_index: "image message"}
                    )

            # Add blank message blocks to the deletion dict
            deletion_dict.update(blank_message_blocks_dict)
            # Add notified text message blocks to the deletion dict
            deletion_dict.update(notified_text_message_blocks_dict)
            # Add deletion image message blocks to the deletion dict
            if len(text_message_blocks_dict) == len(notified_text_message_blocks_dict):
                deletion_dict.update(deletion_images_dict)
            sorted_deletion_dict = dict(sorted(deletion_dict.items()))
            return sorted_deletion_dict

        if len(self.message_blocks_list) != 0:
            sorted_deletion_dict = obtain_blank_and_notified_message_blocks()
            if len(sorted_deletion_dict) != 0:
                # The dict does not fuction well.
                # As processed indexes may be the same.
                deletion_waiting_list = []
                for index, (key, value) in enumerate(sorted_deletion_dict.items()):
                    deletion_waiting_list.append({key - index: value})
                # Feishu server executes the deletion step by step
                for deletion_waiting_dict in deletion_waiting_list:
                    print("deletion_waiting_dict:", deletion_waiting_dict)
                    self.delete_document_children_blocks(
                        list(deletion_waiting_dict.keys())[0]
                    )


if __name__ == "__main__":
    debug_dev_document = True
    now = datetime.now()
    formatted_time = now.strftime("%Y-%m-%d %H:%M:%S")
    print(f"now: {formatted_time}")
    script_dir = Path(__file__).resolve().parent
    configuration_path = script_dir / "configuration.txt"
    with open(configuration_path, "r", encoding="utf-8") as file:
        txt_contents = file.read()
    notice_dir = re.findall(r"notice_dir:\s*(.*)", txt_contents)
    notice_message_heading = re.findall(r"notice_message_heading:\s*(.*)", txt_contents)
    if debug_dev_document is True:
        document_id = re.findall(r"dev_document_id:\s*(.*)", txt_contents)
    else:
        document_id = re.findall(r"document_id:\s*(.*)", txt_contents)
    app_id = re.findall(r"app_id:\s*(.*)", txt_contents)
    app_secret = re.findall(r"app_secret:\s*(.*)", txt_contents)

    feishu_docs_api = FeishuDocsAPI(
        notice_message_heading[0],
        notice_dir[0],
        document_id[0],
        app_id[0],
        app_secret[0],
        message_heading_text="家长留言区",
        save_all_document_blocks_response_as_json=False,
        save_all_document_comments_response_as_json=True,
    )

    # feishu_docs_api.deliver_and_reply_messages()

    feishu_docs_api.delete_notified_messages()
